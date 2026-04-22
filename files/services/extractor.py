"""파일에서 텍스트를 추출하는 모듈.

확장자를 보고 적절한 파서를 골라 순수 텍스트 문자열로 돌려준다.

- .txt, .md : UTF-8(또는 CP949) 그대로 읽기
- .pdf      : 1순위 PyMuPDF(fitz) — 한글 레이아웃 읽기 순서 품질이 가장 좋음.
              실패 시 pdfplumber 폴백(표 마크다운 변환 지원).
              그래도 실패하면 pypdf 최후 폴백.
- .docx     : python-docx로 문단 + 표 모두 추출.
- .xlsx/.xlsm : openpyxl로 시트별 마크다운 변환(병합 셀·수식 캐시값 포함).
- .xls    : xlrd로 레거시 엑셀 시트별 마크다운 변환.

표는 마크다운 형식( | col | col | )으로 변환해서 본문에 끼워넣는다.
LLM이 마크다운 표를 자연스럽게 이해하므로 RAG 검색 품질이 향상된다.
"""

import datetime
import logging
from pathlib import Path
from typing import Any, List

import fitz  # PyMuPDF
import openpyxl
import pdfplumber
import xlrd
from pypdf import PdfReader
from docx import Document as DocxDocument


logger = logging.getLogger(__name__)


class UnsupportedFileType(Exception):
    """지원하지 않는 확장자일 때."""


class EmptyTextError(Exception):
    """파일은 읽었는데 추출된 텍스트가 비어있을 때."""


def extract_text(file_obj, original_name: str) -> str:
    """파일에서 텍스트를 꺼낸다."""
    ext = Path(original_name).suffix.lower()

    if ext in ('.txt', '.md'):
        text = _extract_plain(file_obj)
    elif ext == '.pdf':
        text = _extract_pdf(file_obj)
    elif ext == '.docx':
        text = _extract_docx(file_obj)
    elif ext in ('.xlsx', '.xlsm'):
        text = _extract_xlsx(file_obj)
    elif ext == '.xls':
        text = _extract_xls(file_obj)
    else:
        raise UnsupportedFileType(f'지원하지 않는 확장자: {ext}')

    text = text.strip()
    if not text:
        raise EmptyTextError('파일에서 추출된 텍스트가 없습니다.')
    return text


# ----------------------------------------------------------------------------
# TXT / MD
# ----------------------------------------------------------------------------

def _extract_plain(file_obj) -> str:
    """평문 파일 읽기. UTF-8 우선, 실패 시 CP949로 재시도."""
    file_obj.seek(0)
    raw = file_obj.read()
    try:
        return raw.decode('utf-8')
    except UnicodeDecodeError:
        return raw.decode('cp949', errors='replace')


# ----------------------------------------------------------------------------
# PDF — pdfplumber 우선, 실패 시 pypdf 폴백
# ----------------------------------------------------------------------------

def _extract_pdf(file_obj) -> str:
    """PDF에서 텍스트를 추출. PyMuPDF → pdfplumber → pypdf 순서로 폴백.

    PyMuPDF는 한국어·복잡 레이아웃 PDF에서 읽기 순서를 가장 잘 보존한다.
    표가 많은 PDF가 PyMuPDF에서 실패하면 pdfplumber(마크다운 표)로 폴백.
    """
    try:
        file_obj.seek(0)
        text = _extract_pdf_with_fitz(file_obj)
        if text.strip():
            return text
        raise RuntimeError('fitz returned empty')
    except Exception as e1:
        logger.warning('PyMuPDF 실패, pdfplumber로 폴백: %s', e1)
        try:
            file_obj.seek(0)
            return _extract_pdf_with_plumber(file_obj)
        except Exception as e2:
            logger.warning('pdfplumber 실패, pypdf로 폴백: %s', e2)
            file_obj.seek(0)
            return _extract_pdf_with_pypdf(file_obj)


def _extract_pdf_with_fitz(file_obj) -> str:
    """PyMuPDF로 PDF 텍스트 추출. 읽기 순서(reading order)를 자연스레 복원."""
    data = file_obj.read()
    doc = fitz.open(stream=data, filetype='pdf')
    try:
        pages_text = []
        for page in doc:
            # 'text' 모드: 블록 단위로 읽기 순서대로 텍스트 반환
            text = page.get_text('text') or ''
            if text.strip():
                pages_text.append(text)
        return '\n\n'.join(pages_text)
    finally:
        doc.close()


def _extract_pdf_with_plumber(file_obj) -> str:
    pages_output: List[str] = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            # 1) 먼저 이 페이지의 표를 마크다운으로 변환
            tables = page.extract_tables() or []
            table_markdowns = [_table_to_markdown(t) for t in tables]

            # 2) 페이지 본문 텍스트
            text = page.extract_text() or ''

            # 3) 텍스트 뒤에 표를 이어붙임 (표 위치 보존은 복잡해서 생략)
            parts = []
            if text.strip():
                parts.append(text.strip())
            parts.extend(table_markdowns)

            if parts:
                pages_output.append('\n\n'.join(parts))

    return '\n\n'.join(pages_output)


def _extract_pdf_with_pypdf(file_obj) -> str:
    reader = PdfReader(file_obj)
    pages = [page.extract_text() or '' for page in reader.pages]
    return '\n\n'.join(p for p in pages if p.strip())


# ----------------------------------------------------------------------------
# DOCX — 문단 + 표 추출
# ----------------------------------------------------------------------------

def _extract_docx(file_obj) -> str:
    """DOCX의 문단과 표를 모두 추출.

    python-docx는 문서를 순회할 때 문단·표를 순서대로 방문할 수 있도록
    element.body.iter() 같은 API를 제공하지만 복잡하므로,
    여기서는 단순화를 위해 '문단 전체 → 표 전체' 순서로 이어붙인다.
    (완벽한 원본 순서는 아니지만 대부분 내용은 보존됨)
    """
    file_obj.seek(0)
    doc = DocxDocument(file_obj)

    # 1) 본문 문단
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 2) 표 (각각 마크다운으로)
    table_markdowns: List[str] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append(cells)
        md = _table_to_markdown(rows)
        if md:
            table_markdowns.append(md)

    parts: List[str] = []
    if paragraphs:
        parts.append('\n'.join(paragraphs))
    parts.extend(table_markdowns)

    return '\n\n'.join(parts)


# ----------------------------------------------------------------------------
# XLSX / XLSM — openpyxl (수식 캐시값 + 병합 셀 복제)
# ----------------------------------------------------------------------------

# 시트당 처리할 최대 행 수 (초과분은 경고 + 절단)
XLSX_MAX_ROWS_PER_SHEET = 10_000


def _extract_xlsx(file_obj) -> str:
    """XLSX/XLSM 워크북을 시트별 마크다운 섹션으로 변환."""
    file_obj.seek(0)
    wb = openpyxl.load_workbook(file_obj, data_only=True, read_only=False)
    sections: List[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if ws.sheet_state != 'visible':
            continue

        rows = _xlsx_rows_with_merged(ws)
        truncated = False
        if len(rows) > XLSX_MAX_ROWS_PER_SHEET:
            logger.warning(
                'XLSX 시트 %r: %d행 중 앞 %d행만 추출',
                sheet_name, len(rows), XLSX_MAX_ROWS_PER_SHEET,
            )
            rows = rows[:XLSX_MAX_ROWS_PER_SHEET]
            truncated = True

        # 빈 행 제거 (마지막에 남아있는 완전 빈 행)
        rows = [r for r in rows if any(cell != '' for cell in r)]
        if not rows:
            continue

        body = _rows_to_text(rows)
        if not body:
            continue

        header = f'## {sheet_name}'
        if truncated:
            header += f' (앞 {XLSX_MAX_ROWS_PER_SHEET}행만 포함)'
        sections.append(f'{header}\n\n{body}')

    wb.close()
    return '\n\n'.join(sections)


def _xlsx_rows_with_merged(ws) -> List[List[str]]:
    """openpyxl 워크시트를 2D 문자열 리스트로.

    병합 셀의 경우 좌상단 값만 들어있고 나머지는 None이므로,
    병합 범위 전체에 같은 값을 복제해서 의미 유지.
    """
    # 병합 셀 맵 구성: (row, col) → 좌상단 값
    merged_map: dict = {}
    for merged_range in list(ws.merged_cells.ranges):
        min_row, min_col = merged_range.min_row, merged_range.min_col
        max_row, max_col = merged_range.max_row, merged_range.max_col
        top_left = ws.cell(row=min_row, column=min_col).value
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                merged_map[(r, c)] = top_left

    rows: List[List[str]] = []
    for row_idx, row in enumerate(ws.iter_rows(), start=1):
        cells: List[str] = []
        for cell in row:
            if (cell.row, cell.column) in merged_map:
                value = merged_map[(cell.row, cell.column)]
            else:
                value = cell.value
            cells.append(_format_cell_value(value))
        rows.append(cells)
    return rows


# ----------------------------------------------------------------------------
# XLS — xlrd (레거시 포맷)
# ----------------------------------------------------------------------------

def _extract_xls(file_obj) -> str:
    """XLS (레거시) 워크북을 시트별 마크다운 섹션으로 변환."""
    file_obj.seek(0)
    data = file_obj.read()
    book = xlrd.open_workbook(file_contents=data, formatting_info=True)
    sections: List[str] = []

    for sheet_idx in range(book.nsheets):
        sheet = book.sheet_by_index(sheet_idx)
        if getattr(sheet, 'visibility', 0) != 0:
            # 0 = visible, 1 = hidden, 2 = very hidden
            continue

        # 병합 셀 처리
        merged_map: dict = {}
        for (rlo, rhi, clo, chi) in sheet.merged_cells:
            top_left = sheet.cell(rlo, clo)
            top_left_val = _xls_cell_value(top_left, book)
            for r in range(rlo, rhi):
                for c in range(clo, chi):
                    merged_map[(r, c)] = top_left_val

        n_rows = sheet.nrows
        truncated = False
        if n_rows > XLSX_MAX_ROWS_PER_SHEET:
            logger.warning(
                'XLS 시트 %r: %d행 중 앞 %d행만 추출',
                sheet.name, n_rows, XLSX_MAX_ROWS_PER_SHEET,
            )
            n_rows = XLSX_MAX_ROWS_PER_SHEET
            truncated = True

        rows: List[List[str]] = []
        for r in range(n_rows):
            cells: List[str] = []
            for c in range(sheet.ncols):
                if (r, c) in merged_map:
                    cells.append(_format_cell_value(merged_map[(r, c)]))
                else:
                    cells.append(_format_cell_value(_xls_cell_value(sheet.cell(r, c), book)))
            rows.append(cells)

        rows = [r for r in rows if any(cell != '' for cell in r)]
        if not rows:
            continue

        body = _rows_to_text(rows)
        if not body:
            continue

        header = f'## {sheet.name}'
        if truncated:
            header += f' (앞 {XLSX_MAX_ROWS_PER_SHEET}행만 포함)'
        sections.append(f'{header}\n\n{body}')

    return '\n\n'.join(sections)


def _xls_cell_value(cell, book) -> Any:
    """xlrd 셀의 타입별 원시값을 파이썬 네이티브 타입으로 변환."""
    if cell.ctype == xlrd.XL_CELL_DATE:
        # xldate를 datetime으로
        try:
            return xlrd.xldate.xldate_as_datetime(cell.value, book.datemode)
        except Exception:
            return cell.value
    if cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return bool(cell.value)
    if cell.ctype == xlrd.XL_CELL_EMPTY:
        return None
    return cell.value


# ----------------------------------------------------------------------------
# 셀 값 포맷 공용 (xlsx/xls 공통)
# ----------------------------------------------------------------------------

def _format_cell_value(value) -> str:
    """셀 값을 사람·LLM 친화적인 문자열로 변환."""
    if value is None:
        return ''
    if isinstance(value, datetime.datetime):
        # 시간 정보가 의미 없으면 날짜만
        if value.hour == 0 and value.minute == 0 and value.second == 0:
            return value.strftime('%Y-%m-%d')
        return value.strftime('%Y-%m-%d %H:%M:%S')
    if isinstance(value, datetime.date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float):
        # 정수값이면 정수로 (예: 5.0 → 5)
        if value.is_integer():
            return str(int(value))
        # 소수점은 그대로
        return f'{value:g}'
    return str(value).strip()


# ----------------------------------------------------------------------------
# 표 → 텍스트 공용 변환 (와이드 희소 매트릭스 감지 → key-value 직렬화)
# ----------------------------------------------------------------------------

# 와이드 매트릭스로 간주할 열 수 하한
MATRIX_MIN_COLS = 10
# 와이드 매트릭스로 간주할 빈 셀 비율 하한 (0.0~1.0)
MATRIX_EMPTY_RATIO = 0.6


def _rows_to_text(rows) -> str:
    """2D 셀 배열을 텍스트로. 와이드 희소 매트릭스면 key-value, 아니면 마크다운 표.

    대부분의 표는 기존 마크다운 표로 잘 읽히지만, 결재 라인표·권한 매트릭스
    같은 "열이 많고 데이터 행이 희소한" 표는 파이프 테이블로 표현하면 LLM이
    어느 열이 어느 헤더에 해당하는지 잃어버린다. 이런 경우만 감지해서
    행별 `헤더: 값` 형식으로 직렬화한다.
    """
    if not rows:
        return ''
    col_count = max(len(r) for r in rows)
    normalized = [list(r) + [''] * (col_count - len(r)) for r in rows]

    header_block = _detect_header_block(normalized, col_count)
    if _is_wide_sparse_matrix(normalized, col_count, header_block):
        return _matrix_to_keyvalue(normalized, col_count, header_block)
    return _table_to_markdown(normalized)


def _detect_header_block(rows, col_count: int) -> tuple:
    """헤더 블록 (start, end) 인덱스를 추정.

    - start: 위에서부터 스캔해 채워진 비율 ≥ 70%인 첫 번째 행(헤더는 대개 조밀)
    - end: start 아래로 확장하며 채워진 비율 ≥ 50%인 행까지(다단 헤더 대응).
      데이터 행은 와이드 희소 매트릭스에서 보통 50% 미만이므로 이 임계치로
      헤더와 데이터를 분리.
    그런 행이 없으면 (0, 0) 폴백.
    """
    if not rows or col_count == 0:
        return (0, 0)
    start = None
    for idx, row in enumerate(rows):
        filled = sum(1 for c in row if str(c).strip())
        if filled / col_count >= 0.7:
            start = idx
            break
    if start is None:
        return (0, 0)
    end = start
    for idx in range(start + 1, len(rows)):
        filled = sum(1 for c in rows[idx] if str(c).strip())
        if filled / col_count >= 0.5:
            end = idx
        else:
            break
    return (start, end)


def _is_wide_sparse_matrix(rows, col_count: int, header_block: tuple) -> bool:
    """와이드 희소 매트릭스 여부 판정.

    기준: 열 수 ≥ MATRIX_MIN_COLS AND 데이터 행(헤더 블록 이하)의 빈 셀 비율 ≥
    MATRIX_EMPTY_RATIO. 헤더는 대개 조밀하므로 데이터 행만 본질적 기준.
    """
    if col_count < MATRIX_MIN_COLS:
        return False
    _, header_end = header_block
    data_rows = rows[header_end + 1:]
    if not data_rows:
        return False
    total = col_count * len(data_rows)
    if total == 0:
        return False
    empty = sum(1 for r in data_rows for c in r if not str(c).strip())
    return (empty / total) >= MATRIX_EMPTY_RATIO


def _matrix_to_keyvalue(rows, col_count: int, header_block: tuple) -> str:
    """와이드 희소 매트릭스를 행별 `헤더: 값` 형식으로 직렬화.

    헤더 블록 이전은 타이틀/설명(prelude)으로 앞에 붙이고,
    헤더 블록 각 열 값을 합쳐 다단 헤더를 만든 뒤,
    데이터 행마다 비어있지 않은 열만 `헤더: 값`으로 나열.
    """
    header_start, header_end = header_block

    # 헤더 블록 이전: 타이틀/설명(prelude)
    prelude: List[str] = []
    for i in range(header_start):
        parts = [str(c).strip() for c in rows[i] if str(c).strip()]
        if parts:
            prelude.append(' '.join(parts))

    # 다단 헤더: header_block 범위의 각 열 값을 중복 제거 후 결합
    col_headers: List[str] = []
    for c_idx in range(col_count):
        seen: List[str] = []
        for r_idx in range(header_start, header_end + 1):
            v = str(rows[r_idx][c_idx]).strip()
            if v and v not in seen:
                seen.append(v)
        col_headers.append(' '.join(seen) if seen else f'열{c_idx + 1}')

    # 데이터 행: 비어있지 않은 셀만 `헤더: 값` 페어로
    data_lines: List[str] = []
    seq = 1
    for i in range(header_end + 1, len(rows)):
        pairs = []
        for c_idx, val in enumerate(rows[i]):
            v = str(val).strip()
            if v:
                pairs.append(f'{col_headers[c_idx]}: {v}')
        if pairs:
            data_lines.append(f'[행 {seq}] ' + ' | '.join(pairs))
            seq += 1

    out_parts: List[str] = []
    if prelude:
        out_parts.append('\n'.join(prelude))
    if data_lines:
        out_parts.append('\n'.join(data_lines))
    return '\n\n'.join(out_parts)


def _table_to_markdown(rows) -> str:
    """2차원 리스트(행 x 셀)를 마크다운 표로 변환.

    빈 셀은 '-'로 치환, 빈 표는 빈 문자열 반환.
    """
    if not rows:
        return ''

    # 셀 정리 (None, 개행 등)
    def clean(cell):
        if cell is None:
            return ''
        return str(cell).replace('\n', ' ').replace('|', '/').strip()

    cleaned = [[clean(c) for c in row] for row in rows]
    cleaned = [r for r in cleaned if any(cell for cell in r)]  # 빈 행 제거
    if not cleaned:
        return ''

    col_count = max(len(r) for r in cleaned)
    # 열 개수 맞추기
    normalized = [r + [''] * (col_count - len(r)) for r in cleaned]

    # 마크다운 행 조립
    lines = []
    header = normalized[0]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('|' + '|'.join([' --- '] * col_count) + '|')
    for row in normalized[1:]:
        lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(lines)
